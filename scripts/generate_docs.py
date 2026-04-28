"""
Generador del documento Word con la documentacion completa del proyecto
Servicio EUROPCAR V1.

Salida: Contexto/Documentacion_Servicio_Europcar_V1.docx

Este script construye el .docx con python-docx. No depende de plantillas externas
y se puede regenerar tras cambios en codigo/BD ajustando las tablas/datos abajo.
"""

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = REPO_ROOT / "Contexto" / "Documentacion_Servicio_Europcar_V1.docx"
IMG_DIR = REPO_ROOT / "Contexto" / "imagenes_origen"
SEQ_DIR = REPO_ROOT / "Contexto" / "diagramas_secuencia"


def img(name: str) -> Path:
    """Resuelve la ruta de una figura extraida del docx original."""
    return IMG_DIR / name


def seq(name: str) -> Path:
    """Resuelve la ruta de un diagrama de secuencia generado."""
    return SEQ_DIR / name


# ---------------------------------------------------------------------------
# Helpers de estilo
# ---------------------------------------------------------------------------

PRIMARY = RGBColor(0x0F, 0x4C, 0x81)   # azul corporativo
ACCENT = RGBColor(0x2E, 0x86, 0xC1)
DARK_TEXT = RGBColor(0x1B, 0x26, 0x31)
SOFT_TEXT = RGBColor(0x4A, 0x55, 0x68)


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex: str = "BFC9D1", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False,
                  size: int = 11, color: RGBColor | None = None,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = color or DARK_TEXT
    p.paragraph_format.space_after = Pt(space_after)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(0.6 * level)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_TEXT


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              header_fill: str = "0F4C81", widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light List Accent 1"

    hdr = table.rows[0].cells
    for i, name in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(name)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        _set_cell_shading(hdr[i], header_fill)
        _set_cell_border(hdr[i], color_hex="0F4C81", size="6")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            run = para.add_run(str(value) if value is not None else "")
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_TEXT
            _set_cell_border(cells[i])

    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)


def add_image(doc: Document, path: Path, caption: str | None = None,
              width_cm: float = 15.0) -> None:
    """Inserta una imagen centrada con caption opcional debajo."""
    if not path.exists():
        add_paragraph(doc, f"[Imagen no encontrada: {path.name}]",
                      italic=True, color=SOFT_TEXT,
                      align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(path), width=Cm(width_cm))
    except Exception as exc:  # imagen corrupta o formato no soportado
        add_paragraph(doc, f"[No se pudo insertar {path.name}: {exc}]",
                      italic=True, color=SOFT_TEXT,
                      align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        return
    if caption:
        add_paragraph(doc, caption, italic=True, size=10, color=SOFT_TEXT,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F4C81")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Datos del documento
# ---------------------------------------------------------------------------

ENDPOINTS_AUTH = [
    ["POST", "/api/v1/Auth/login", "Inicio de sesion. Devuelve JWT, roles y datos de usuario.", "Anonimo"],
    ["POST", "/api/v1/Auth/register", "Registro de cliente y creacion de usuario asociado.", "Anonimo"],
    ["GET",  "/api/v1/Auth/cedula-exists?cedula=...", "Verifica si una identificacion ya existe.", "Anonimo"],
]

ENDPOINTS_BOOKING_VEH = [
    ["GET",  "/api/v1/vehiculos", "Busqueda paginada de vehiculos disponibles para integracion (Booking).", "Publico"],
    ["GET",  "/api/v1/vehiculos/{vehiculoId}", "Detalle de un vehiculo por codigo interno alfanumerico.", "Publico"],
    ["GET",  "/api/v1/vehiculos/{vehiculoId}/disponibilidad", "Disponibilidad en tiempo real para fechas y localizacion.", "Publico"],
]

ENDPOINTS_BOOKING_CAT = [
    ["GET",  "/api/v1/localizaciones", "Listado paginado de sucursales (con filtros opcionales).", "Publico"],
    ["GET",  "/api/v1/localizaciones/{localizacionId}", "Detalle de localizacion.", "Publico"],
    ["GET",  "/api/v1/categorias", "Catalogo de categorias de vehiculos.", "Publico"],
    ["GET",  "/api/v1/extras", "Catalogo de extras y precios.", "Publico"],
]

ENDPOINTS_BOOKING_RES = [
    ["POST",  "/api/v1/reservas", "Crear una nueva reserva (canal externo).", "Publico"],
    ["GET",   "/api/v1/reservas/{codigoReserva}", "Detalle de una reserva por codigo.", "Publico"],
    ["PATCH", "/api/v1/reservas/{codigoReserva}/cancelar", "Cancelar reserva externa.", "Publico"],
    ["GET",   "/api/v1/reservas/{codigoReserva}/factura", "Factura asociada a la reserva.", "Publico"],
]

ENDPOINTS_ADMIN_VEHICULOS = [
    ["GET",    "/api/v1/admin/vehiculos", "Listado de la flota activa.", "ADMIN, AGENTE_POS"],
    ["GET",    "/api/v1/admin/vehiculos/disponibles", "Busqueda con filtros para back-office.", "Publico"],
    ["GET",    "/api/v1/admin/vehiculos/{id}", "Detalle de un vehiculo por id.", "Publico"],
    ["POST",   "/api/v1/admin/vehiculos", "Alta de vehiculo en la flota.", "ADMIN, AGENTE_POS"],
    ["PUT",    "/api/v1/admin/vehiculos/{id}", "Edicion de vehiculo (incluye imagen).", "ADMIN, AGENTE_POS"],
    ["DELETE", "/api/v1/admin/vehiculos/{id}", "Soft delete de vehiculo.", "ADMIN"],
]

ENDPOINTS_ADMIN_LOCALIZ = [
    ["GET",    "/api/v1/admin/localizaciones?soloActivas=", "Listado de sucursales (incluye inactivas).", "ADMIN, AGENTE_POS"],
    ["GET",    "/api/v1/admin/localizaciones/{id}", "Detalle de sucursal.", "ADMIN, AGENTE_POS"],
    ["GET",    "/api/v1/admin/localizaciones/ciudades", "Ciudades para selectores.", "ADMIN, AGENTE_POS"],
    ["POST",   "/api/v1/admin/localizaciones", "Crear sucursal.", "ADMIN"],
    ["PUT",    "/api/v1/admin/localizaciones/{id}", "Actualizar sucursal.", "ADMIN, AGENTE_POS"],
    ["PUT",    "/api/v1/admin/localizaciones/{id}/estado", "Activar/inhabilitar.", "ADMIN"],
    ["DELETE", "/api/v1/admin/localizaciones/{id}", "Soft delete.", "ADMIN"],
]

ENDPOINTS_ADMIN_RESERVAS = [
    ["POST", "/api/v1/admin/reservas/guest-client", "Crea/busca cliente invitado para flujo de reserva.", "Anonimo"],
    ["POST", "/api/v1/admin/reservas", "Crear una reserva interna.", "Autenticado"],
    ["GET",  "/api/v1/admin/reservas/{codigo}", "Obtener reserva por codigo.", "Autenticado"],
    ["GET",  "/api/v1/admin/reservas/cliente/{idCliente}", "Reservas de un cliente.", "Autenticado (mismo cliente o ADMIN/AGENTE_POS)"],
    ["PUT",  "/api/v1/admin/reservas/{id}/confirmar", "Confirmar reserva + pago + factura.", "Autenticado"],
    ["PUT",  "/api/v1/admin/reservas/{id}/cancelar", "Cancelar reserva (con motivo).", "Autenticado"],
]

ENDPOINTS_ADMIN_CONTRATOS = [
    ["GET",  "/api/v1/Contratos", "Listado de contratos.", "ADMIN, AGENTE_POS"],
    ["GET",  "/api/v1/Contratos/{id}", "Detalle de contrato.", "ADMIN, AGENTE_POS"],
    ["POST", "/api/v1/Contratos", "Crear contrato desde reserva confirmada.", "ADMIN, AGENTE_POS"],
    ["POST", "/api/v1/Contratos/checkout", "Registrar entrega del vehiculo (check-out).", "ADMIN, AGENTE_POS"],
    ["POST", "/api/v1/Contratos/checkin", "Registrar devolucion (check-in).", "ADMIN, AGENTE_POS"],
]

ENDPOINTS_ADMIN_PAGOS = [
    ["GET",  "/api/v1/Pagos", "Listado de pagos.", "ADMIN, AGENTE_POS"],
    ["GET",  "/api/v1/Pagos/{id}", "Detalle de pago.", "ADMIN, AGENTE_POS"],
    ["GET",  "/api/v1/Pagos/reserva/{idReserva}", "Pagos por reserva.", "ADMIN, AGENTE_POS"],
    ["POST", "/api/v1/Pagos", "Registrar pago.", "Autenticado"],
]

ENDPOINTS_ADMIN_FACTURAS = [
    ["GET", "/api/v1/Facturas/mis-facturas", "Facturas del cliente autenticado.", "CLIENTE_WEB"],
]

ENDPOINTS_ADMIN_CATALOGOS = [
    ["GET", "/api/v1/Catalogos/localizaciones", "Sucursales activas.", "Autenticado"],
    ["GET", "/api/v1/Catalogos/localizaciones/{id}", "Detalle de sucursal.", "Autenticado"],
    ["GET", "/api/v1/Catalogos/categorias", "Categorias de vehiculos.", "Autenticado"],
    ["GET", "/api/v1/Catalogos/marcas", "Marcas de vehiculos.", "Autenticado"],
    ["GET", "/api/v1/Catalogos/extras", "Extras y precios.", "Autenticado"],
]

ENDPOINTS_ADMIN_CLIENTES = [
    ["GET",    "/api/v1/Clientes", "Listado de clientes.", "ADMIN, AGENTE_POS"],
    ["GET",    "/api/v1/Clientes/{id}", "Detalle de cliente.", "ADMIN, AGENTE_POS"],
    ["POST",   "/api/v1/Clientes", "Alta de cliente.", "ADMIN, AGENTE_POS"],
    ["PUT",    "/api/v1/Clientes/{id}", "Actualizar cliente.", "ADMIN, AGENTE_POS"],
    ["DELETE", "/api/v1/Clientes/{id}", "Soft delete.", "ADMIN"],
]

ENDPOINTS_ADMIN_USUARIOS = [
    ["GET",    "/api/v1/Usuarios", "Listado de usuarios y sus roles.", "ADMIN"],
    ["GET",    "/api/v1/Usuarios/{id}", "Detalle de usuario.", "ADMIN"],
    ["POST",   "/api/v1/Usuarios", "Crear usuario con cliente asociado y roles.", "ADMIN"],
    ["PUT",    "/api/v1/Usuarios/{id}/estado", "Cambiar estado (ACT/INACT).", "ADMIN"],
    ["PUT",    "/api/v1/Usuarios/{id}/roles", "Actualizar roles asignados.", "ADMIN"],
    ["DELETE", "/api/v1/Usuarios/{id}", "Soft delete.", "ADMIN"],
]

ENDPOINTS_ADMIN_MANTENIMIENTOS = [
    ["GET",  "/api/v1/Mantenimientos", "Listado de mantenimientos.", "ADMIN, AGENTE_POS"],
    ["GET",  "/api/v1/Mantenimientos/{id}", "Detalle de mantenimiento.", "ADMIN, AGENTE_POS"],
    ["GET",  "/api/v1/Mantenimientos/vehiculo/{idVehiculo}", "Mantenimientos de un vehiculo.", "ADMIN, AGENTE_POS"],
    ["POST", "/api/v1/Mantenimientos", "Enviar vehiculo a taller.", "ADMIN, AGENTE_POS"],
    ["PUT",  "/api/v1/Mantenimientos/{id}/cerrar", "Cerrar mantenimiento.", "ADMIN, AGENTE_POS"],
]

ENDPOINTS_HEALTH = [
    ["GET", "/health/live", "Liveness del proceso.", "Publico"],
    ["GET", "/health/ready", "Readiness (incluye chequeo de BD).", "Publico"],
    ["GET", "/swagger", "Documentacion interactiva OpenAPI.", "Publico"],
]


USE_CASES = [
    {
        "id": "CU-01", "nombre": "Iniciar sesion",
        "actor": "Usuario interno (ADMIN/AGENTE_POS) o Cliente web",
        "precondiciones": "Usuario registrado y activo en el sistema.",
        "flujo": [
            "El usuario abre la aplicacion y selecciona Iniciar sesion.",
            "Ingresa username/correo y contrasena.",
            "El sistema valida credenciales contra la API (POST /api/v1/Auth/login).",
            "La API responde con JWT, roles y datos del usuario.",
            "El frontend redirige al dashboard correspondiente segun el tipo de usuario.",
        ],
        "alternativos": "Credenciales invalidas: se notifica error sin revelar detalles. Cuenta bloqueada/inactiva: rechazo controlado.",
        "postcondiciones": "Sesion activa con token JWT y rol asignado.",
    },
    {
        "id": "CU-02", "nombre": "Registrar cliente nuevo",
        "actor": "Visitante del sitio web",
        "precondiciones": "Identificacion no registrada previamente.",
        "flujo": [
            "El visitante accede a /registro.",
            "Completa el formulario (nombres, identificacion, correo, contrasena).",
            "El frontend valida formato (Zod + react-hook-form) y verifica cedula con GET /api/v1/Auth/cedula-exists.",
            "Se envia POST /api/v1/Auth/register; la API crea el cliente y el usuario CLIENTE_WEB.",
            "El sistema responde con login automatico opcional.",
        ],
        "alternativos": "Cedula ya existente: bloqueo y mensaje. Datos invalidos: errores por campo.",
        "postcondiciones": "Cliente y usuario activos en el sistema.",
    },
    {
        "id": "CU-03", "nombre": "Buscar y reservar vehiculo",
        "actor": "Cliente web (autenticado o invitado)",
        "precondiciones": "Existen sucursales y vehiculos activos en la flota.",
        "flujo": [
            "El cliente abre la pagina de inicio o /buscar e ingresa fechas, localizacion de retiro y devolucion.",
            "El frontend consulta GET /api/v1/vehiculos con los filtros.",
            "El cliente revisa el detalle de un vehiculo (GET /api/v1/vehiculos/{id}).",
            "Confirma fechas y entra al flujo /reservar/:id.",
            "Si es invitado, se ejecuta POST /api/v1/admin/reservas/guest-client para localizar/crear el cliente.",
            "Se crea la reserva con POST /api/v1/admin/reservas o POST /api/v1/reservas.",
            "El sistema retorna el codigo de reserva y resumen.",
        ],
        "alternativos": "Vehiculo deja de estar disponible (reserva concurrente): mensaje y reseleccion. Lead time minimo no cumplido: rechazo.",
        "postcondiciones": "Reserva creada en estado PENDIENTE.",
    },
    {
        "id": "CU-04", "nombre": "Confirmar reserva y registrar pago",
        "actor": "Agente POS / Cliente",
        "precondiciones": "Reserva en estado PENDIENTE.",
        "flujo": [
            "El agente abre la reserva en el back-office.",
            "Captura medio de pago y monto.",
            "Se invoca PUT /api/v1/admin/reservas/{id}/confirmar.",
            "La API genera Pago y Factura y cambia el estado de la reserva a CONFIRMADA.",
        ],
        "alternativos": "Pago rechazado: la reserva permanece PENDIENTE.",
        "postcondiciones": "Reserva CONFIRMADA con pago y factura asociados.",
    },
    {
        "id": "CU-05", "nombre": "Cancelar reserva",
        "actor": "Cliente o ADMIN/AGENTE_POS",
        "precondiciones": "Reserva confirmada o pendiente; fecha de retiro futura para clientes.",
        "flujo": [
            "El usuario selecciona Cancelar e ingresa motivo.",
            "El sistema valida propiedad y vigencia (PUT /api/v1/admin/reservas/{id}/cancelar).",
            "La API anula pagos y facturas asociados y cambia estado a CANCELADA.",
        ],
        "alternativos": "Reserva pasada o ajena: error 403.",
        "postcondiciones": "Reserva CANCELADA con auditoria.",
    },
    {
        "id": "CU-06", "nombre": "Apertura de contrato (check-out)",
        "actor": "AGENTE_POS / ADMIN",
        "precondiciones": "Reserva CONFIRMADA y cliente presente en sucursal.",
        "flujo": [
            "El agente crea el contrato a partir de la reserva (POST /api/v1/Contratos).",
            "Registra check-out con kilometraje y nivel de combustible (POST /api/v1/Contratos/checkout).",
            "El vehiculo cambia a estado ALQUILADO.",
        ],
        "alternativos": "Cliente sin licencia vigente: bloqueo.",
        "postcondiciones": "Contrato ABIERTO y vehiculo entregado.",
    },
    {
        "id": "CU-07", "nombre": "Devolucion (check-in)",
        "actor": "AGENTE_POS / ADMIN",
        "precondiciones": "Contrato ABIERTO.",
        "flujo": [
            "El agente registra POST /api/v1/Contratos/checkin con kilometraje, combustible, limpieza y observaciones.",
            "La API calcula cargos por combustible, km extra y limpieza.",
            "El contrato pasa a CERRADO; el vehiculo regresa a DISPONIBLE.",
        ],
        "alternativos": "Daños: registro de cargos adicionales.",
        "postcondiciones": "Contrato CERRADO; flota actualizada.",
    },
    {
        "id": "CU-08", "nombre": "Gestionar flota de vehiculos",
        "actor": "ADMIN / AGENTE_POS",
        "precondiciones": "Sesion con rol valido.",
        "flujo": [
            "Listar vehiculos (/api/v1/admin/vehiculos).",
            "Crear, actualizar o eliminar (CRUD).",
        ],
        "alternativos": "Eliminar vehiculo con reservas: bloqueo logico.",
        "postcondiciones": "Flota actualizada.",
    },
    {
        "id": "CU-09", "nombre": "Registrar mantenimiento",
        "actor": "ADMIN / AGENTE_POS",
        "precondiciones": "Vehiculo en estado DISPONIBLE.",
        "flujo": [
            "Crear mantenimiento (POST /api/v1/Mantenimientos).",
            "El vehiculo pasa a TALLER.",
            "Cerrar (PUT /api/v1/Mantenimientos/{id}/cerrar) lo regresa a DISPONIBLE.",
        ],
        "alternativos": "Cierre con costos: actualizacion de historial.",
        "postcondiciones": "Disponibilidad de flota actualizada.",
    },
    {
        "id": "CU-10", "nombre": "Administrar localizaciones",
        "actor": "ADMIN / AGENTE_POS",
        "precondiciones": "Sesion con rol valido.",
        "flujo": [
            "CRUD sobre /api/v1/admin/localizaciones.",
            "Activar/inhabilitar con /estado.",
        ],
        "alternativos": "Eliminar con vehiculos asignados: validacion de negocio.",
        "postcondiciones": "Catalogo de sucursales actualizado.",
    },
    {
        "id": "CU-11", "nombre": "Administrar usuarios y roles",
        "actor": "ADMIN",
        "precondiciones": "Rol ADMIN.",
        "flujo": [
            "Crear usuario (POST /api/v1/Usuarios) que internamente crea cliente y asigna roles.",
            "Cambiar estado o roles segun necesidad operativa.",
        ],
        "alternativos": "Conflicto por username/correo: respuesta 409.",
        "postcondiciones": "Usuarios y permisos actualizados.",
    },
    {
        "id": "CU-12", "nombre": "Consultar mis facturas / reservas / contratos",
        "actor": "CLIENTE_WEB",
        "precondiciones": "Cliente autenticado.",
        "flujo": [
            "Acceder al portal del cliente (/mi-cuenta, /mis-reservas, /mis-contratos, /mis-facturas, /historial).",
            "El frontend consume endpoints filtrados por id de cliente del token.",
        ],
        "alternativos": "Sin registros: estado vacio.",
        "postcondiciones": "Visualizacion segura de informacion personal.",
    },
    {
        "id": "CU-13", "nombre": "Integracion Booking/OTA (busqueda y reserva)",
        "actor": "Sistema externo (Booking)",
        "precondiciones": "Acceso a endpoints publicos /api/v1/...",
        "flujo": [
            "Booking consulta vehiculos, localizaciones, categorias y extras.",
            "Crea/cancela reservas y consulta facturas via endpoints publicos.",
        ],
        "alternativos": "Reglas para evitar overbooking y prioridad del E-commerce.",
        "postcondiciones": "Sincronizacion de inventario y reservas.",
    },
]


# ---------------------------------------------------------------------------
# Construccion del documento
# ---------------------------------------------------------------------------

def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_TEXT

    # ============ Portada ============
    add_image(doc, img("figura_01.png"), width_cm=4.5)
    add_paragraph(doc, "Servicio EUROPCAR V1", bold=True, size=28,
                  color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_paragraph(doc, "Documentacion tecnica del sistema",
                  italic=True, size=16, color=ACCENT,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_paragraph(doc, "Plataforma de renta de vehiculos: backend ASP.NET Core, frontend React + Vite, "
                       "base de datos relacional con esquemas rental, security y audit.",
                  size=12, color=SOFT_TEXT,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_paragraph(doc, "Pontificia Universidad Catolica del Ecuador",
                  bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Facultad de Habitat, Infraestructura y Creatividad — Ingenieria de Software",
                  size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Equipo: Domenica Arcos, Dana Bahamonde, Dylan Medina, Juan Morales",
                  size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_divider(doc)

    # ============ Indice (manual) ============
    add_heading(doc, "Indice", level=1)
    indices = [
        "1. Introduccion y vision general",
        "2. Glosario",
        "3. Estructura del sistema",
        "4. Arquitectura logica y tecnologica",
        "5. Casos de uso",
        "6. Estructura de la base de datos",
        "7. Backend y endpoints de la API",
        "8. Construccion del frontend",
        "9. Seguridad, roles y autorizacion",
        "10. Configuracion, ejecucion y despliegue",
        "11. Requerimientos no funcionales",
        "12. Roadmap y mejoras propuestas",
    ]
    for item in indices:
        add_bullet(doc, item)

    doc.add_page_break()

    # ============ 1. Introduccion ============
    add_heading(doc, "1. Introduccion y vision general", level=1)
    add_paragraph(doc,
        "Servicio EUROPCAR V1 es una plataforma integral de gestion de renta de vehiculos. "
        "Su objetivo es operar el ciclo de vida completo del negocio: catalogo de flota, "
        "busqueda y disponibilidad en tiempo real, motor de reservas, contratos, pagos, "
        "facturacion, mantenimientos y administracion de sucursales y usuarios. La "
        "plataforma se compone de un backend en ASP.NET Core que expone una API REST "
        "versionada, un frontend SPA en React + Vite y una base de datos relacional con "
        "esquemas separados para datos operativos (rental), seguridad (security) y auditoria "
        "(audit). Adicionalmente, el sistema expone endpoints publicos para integracion con "
        "agregadores externos tipo OTA (por ejemplo, Booking).")
    add_paragraph(doc,
        "Esta documentacion detalla la estructura del proyecto, los casos de uso del sistema, "
        "la estructura de la base de datos, el conjunto completo de endpoints del backend y la "
        "construccion del frontend, con foco en proporcionar una vision unica y consolidada del "
        "diseno actual.")

    # ============ 2. Glosario ============
    add_heading(doc, "2. Glosario", level=1)
    glosario = [
        ["Flota", "Conjunto de vehiculos que la empresa posee y gestiona para renta."],
        ["Categoria de vehiculo (SIPP)", "Clasificacion estandar (Economy, Compact, SUV, etc.) usada para reservar."],
        ["Localizacion / Sucursal", "Punto fisico de retiro y entrega del vehiculo."],
        ["Booking engine", "Logica que valida disponibilidad, calcula precio y bloquea inventario."],
        ["Extras / Add-ons", "Servicios adicionales (GPS, silla de bebe, conductor adicional, etc.)."],
        ["Voucher", "Comprobante digital o impreso de la reserva."],
        ["Lead time", "Tiempo minimo de anticipacion para reservar."],
        ["Channel manager", "Componente responsable de mantener disponibilidad y tarifas sincronizadas con OTA."],
        ["Overbooking", "Situacion critica de venta de mas autos de los disponibles; el sistema debe evitarla."],
        ["Rate plan", "Conjunto de reglas aplicadas al precio (no reembolsable, flexible, etc.)."],
        ["JWT", "JSON Web Token usado para autenticar peticiones a la API."],
        ["RBAC", "Modelo de autorizacion basado en roles (ADMIN, AGENTE_POS, CLIENTE_WEB)."],
        ["Soft delete", "Borrado logico mediante un flag (no se elimina la fila)."],
        ["Auditoria", "Registro inmutable de eventos y cambios sobre tablas criticas."],
    ]
    add_table(doc, ["Termino", "Definicion"], glosario, widths_cm=[5.0, 11.5])

    # ============ 3. Estructura del sistema ============
    add_heading(doc, "3. Estructura del sistema", level=1)
    add_paragraph(doc, "El repositorio se organiza en dos grandes piezas: backend (.NET) en src/ y frontend (React) en frontend/. Tambien incluye material de contexto, scripts y configuracion de despliegue.")
    estructura = [
        ["src/Europcar.Rental.Api",
         "Capa de presentacion (.NET 10). Controllers REST, middleware, configuracion de Swagger, JWT, versionado y health checks."],
        ["src/Europcar.Rental.Business",
         "Capa de negocio: servicios, validadores, DTOs, mappers y excepciones del dominio."],
        ["src/Europcar.Rental.DataAccess",
         "Acceso a datos con EF Core: DbContext, entidades por modulo, configuraciones y repositorios."],
        ["src/Europcar.Rental.DataManagement",
         "Servicios de datos auxiliares (modelos planos y operaciones de soporte)."],
        ["frontend",
         "SPA en React + Vite con paginas publicas, panel administrativo y portal del cliente."],
        ["scripts",
         "Utilitarios de soporte (por ejemplo, generador de documentacion)."],
        ["Contexto",
         "Material funcional y academico (incluye EUROPCAR.docx y este documento generado)."],
        [".github / Dockerfile / render.yaml",
         "Automatizacion (CI), contenedor base y blueprint de despliegue del frontend."],
    ]
    add_table(doc, ["Carpeta / Proyecto", "Responsabilidad"], estructura, widths_cm=[6.5, 10.0])

    # ============ 4. Arquitectura ============
    add_heading(doc, "4. Arquitectura logica y tecnologica", level=1)
    add_paragraph(doc,
        "El backend sigue una arquitectura por capas con separacion clara de responsabilidades. "
        "La API monolitica esta dividida en proyectos para facilitar mantenibilidad y una eventual "
        "evolucion a microservicios. La logica de negocio reside en services y la persistencia en EF Core. "
        "El frontend es una SPA desacoplada que consume la API mediante un cliente HTTP centralizado y maneja "
        "estado de servidor con React Query y estado global con Zustand.")

    add_heading(doc, "4.1. Diagrama de arquitectura", level=2)
    add_image(doc, img("figura_17.jpeg"),
              caption="Figura 1. Diagrama de Arquitectura del sistema.",
              width_cm=15.5)

    add_heading(doc, "4.2. Capas del backend", level=2)
    capas = [
        ["Presentacion (Api)", "Controllers, autenticacion JWT, autorizacion por rol, versionado, Swagger, manejo global de errores y health checks."],
        ["Aplicacion / Negocio (Business)", "Servicios de dominio (auth, reservas, contratos, vehiculos, etc.), validacion, mapeo y excepciones."],
        ["Datos (DataAccess)", "DbContext y entidades EF Core para esquemas rental, security y audit."],
        ["Datos auxiliares (DataManagement)", "Modelos y servicios de datos planos para operaciones especificas (clientes, usuarios, facturas)."],
    ]
    add_table(doc, ["Capa", "Responsabilidad"], capas, widths_cm=[5.5, 11.0])

    add_heading(doc, "4.3. Stack tecnologico", level=2)
    stack = [
        ["Lenguaje backend", "C# .NET 10"],
        ["Framework API", "ASP.NET Core (Web API)"],
        ["ORM", "Entity Framework Core"],
        ["Versionado API", "Asp.Versioning.Mvc"],
        ["Documentacion API", "Swashbuckle (Swagger UI)"],
        ["Autenticacion", "JWT (Microsoft.AspNetCore.Authentication.JwtBearer)"],
        ["Health checks", "Microsoft.Extensions.Diagnostics.HealthChecks (incluye EFCore)"],
        ["Base de datos", "PostgreSQL (configuracion para Supabase)"],
        ["Frontend", "React 19 + Vite 8"],
        ["Routing", "React Router 7"],
        ["Estado servidor", "TanStack React Query 5"],
        ["Estado global", "Zustand"],
        ["Formularios", "React Hook Form + Zod"],
        ["HTTP", "Axios con interceptors"],
        ["Notificaciones", "Sonner"],
        ["Iconos", "Lucide React"],
        ["Despliegue front", "Render (blueprint render.yaml)"],
    ]
    add_table(doc, ["Componente", "Tecnologia"], stack, widths_cm=[5.0, 11.5])

    # ============ 5. Casos de uso ============
    add_heading(doc, "5. Casos de uso", level=1)
    add_paragraph(doc,
        "Se describen los casos de uso principales que cubren el flujo del negocio (catalogo, "
        "reserva, pago, contrato, devolucion, mantenimientos) y la administracion del sistema "
        "(usuarios, sucursales, flota), junto con la integracion publica para canales tipo OTA.")

    add_heading(doc, "5.1. Diagrama conceptual", level=2)
    add_image(doc, img("figura_02.png"),
              caption="Figura 2. Diagrama conceptual del proyecto.",
              width_cm=15.5)

    add_heading(doc, "5.2. Diagrama general de funciones", level=2)
    add_image(doc, img("figura_03.png"),
              caption="Figura 3. Diagrama general de funciones del sistema.",
              width_cm=15.5)

    add_heading(doc, "5.3. Diagramas de secuencia de las actividades principales", level=2)
    add_paragraph(doc,
        "Los siguientes diagramas describen, paso a paso, la interaccion entre actores, frontend, "
        "API y base de datos para las cinco actividades criticas del sistema: autenticacion, "
        "busqueda y reserva, confirmacion con pago y factura, apertura del contrato (check-out) y "
        "devolucion del vehiculo (check-in).")

    secuencias = [
        ("seq_01_cu_01_iniciar_sesion.png",
         "Figura 5.1. Diagrama de secuencia: CU-01 Iniciar sesion."),
        ("seq_02_cu_03_buscar_y_reservar_vehiculo.png",
         "Figura 5.2. Diagrama de secuencia: CU-03 Buscar y reservar vehiculo."),
        ("seq_03_cu_04_confirmar_reserva_con_pago_y_factura.png",
         "Figura 5.3. Diagrama de secuencia: CU-04 Confirmar reserva con pago y factura."),
        ("seq_04_cu_06_apertura_de_contrato_check_out.png",
         "Figura 5.4. Diagrama de secuencia: CU-06 Apertura de contrato (check-out)."),
        ("seq_05_cu_07_devolucion_check_in.png",
         "Figura 5.5. Diagrama de secuencia: CU-07 Devolucion (check-in)."),
    ]
    for fname, caption in secuencias:
        add_image(doc, seq(fname), caption=caption, width_cm=16.0)

    add_heading(doc, "5.4. Detalle de casos de uso", level=2)
    for cu in USE_CASES:
        add_heading(doc, f"{cu['id']} — {cu['nombre']}", level=2)
        add_paragraph(doc, f"Actor: {cu['actor']}", bold=True)
        add_paragraph(doc, f"Precondiciones: {cu['precondiciones']}")
        add_paragraph(doc, "Flujo principal:", bold=True)
        for paso in cu["flujo"]:
            add_bullet(doc, paso)
        add_paragraph(doc, f"Flujos alternativos: {cu['alternativos']}")
        add_paragraph(doc, f"Postcondiciones: {cu['postcondiciones']}")

    doc.add_page_break()

    # ============ 6. Base de datos ============
    add_heading(doc, "6. Estructura de la base de datos", level=1)
    add_paragraph(doc,
        "La base de datos se organiza en tres esquemas principales:")
    add_bullet(doc, "rental: dominio operativo del negocio (flota, clientes, reservas, contratos, pagos, facturas, mantenimientos, geografia y catalogos).")
    add_bullet(doc, "security: gestion de usuarios, roles, permisos y sesiones.")
    add_bullet(doc, "audit: registro de eventos del sistema e intentos de inicio de sesion.")
    add_paragraph(doc,
        "Casi todas las entidades operativas heredan de una clase base con auditoria comun "
        "(EsEliminado, FechaRegistroUtc, CreadoPorUsuario, ModificadoPorUsuario, FechaModificacionUtc, "
        "ModificadoDesdeIp, RowVersion). Las que tienen estado adicionalmente incluyen OrigenRegistro, "
        "FechaInhabilitacionUtc y MotivoInhabilitacion.")

    add_heading(doc, "6.0. Diagrama Entidad-Relacion", level=2)
    add_image(doc, img("figura_18.png"),
              caption="Figura 4. Diagrama Entidad-Relacion de la base de datos del sistema.",
              width_cm=16.0)

    # 6.1 Modulo de geografia y catalogos
    add_heading(doc, "6.1. Modulo de geografia y catalogos", level=2)
    add_table(doc, ["Tabla", "Campos clave", "Descripcion"], [
        ["paises", "id_pais (PK), pais_guid, codigo_iso2, nombre_pais, estado_pais", "Catalogo de paises."],
        ["ciudades", "id_ciudad (PK), id_pais (FK), nombre_ciudad, estado_ciudad", "Ciudades por pais."],
        ["localizaciones", "id_localizacion (PK), codigo_localizacion, id_ciudad (FK), direccion, telefono, correo, horario, zona_horaria, lat, lng, estado", "Sucursales fisicas."],
        ["marcas_vehiculos", "id_marca (PK), codigo_marca, nombre_marca, descripcion, estado", "Marcas de vehiculos."],
        ["categorias_vehiculos", "id_categoria (PK), codigo_categoria, nombre_categoria, kilometraje_ilimitado, limite_km_dia, cargo_km_excedente, estado", "Clasificacion SIPP/comercial de vehiculos."],
        ["extras", "id_extra (PK), codigo_extra, nombre, descripcion, tipo_extra, requiere_stock, valor_fijo, estado", "Servicios adicionales (GPS, sillas, etc.)."],
        ["localizacion_extra_stock", "id_loc_extra_stock (PK), id_localizacion (FK), id_extra (FK), stock_disponible, stock_reservado, estado", "Inventario de extras por sucursal."],
    ], widths_cm=[3.6, 6.8, 6.0])

    # 6.2 Modulo flota
    add_heading(doc, "6.2. Modulo de flota", level=2)
    add_table(doc, ["Tabla", "Campos clave", "Descripcion"], [
        ["vehiculos",
         "id_vehiculo (PK), vehiculo_guid, codigo_interno, placa, id_marca (FK), id_categoria (FK), "
         "modelo, anio, color, combustible, transmision, capacidad_pasajeros, capacidad_maletas, "
         "puertas, localizacion_actual (FK), precio_base_dia, kilometraje_actual, aire_acondicionado, "
         "estado_operativo (DISPONIBLE/ALQUILADO/TALLER), imagen_url, estado",
         "Unidades fisicas de la flota."],
        ["mantenimientos",
         "id_mantenimiento (PK), codigo, id_vehiculo (FK), tipo, fecha_inicio, fecha_fin, kilometraje, "
         "costo, proveedor_taller, estado (ABIERTO/CERRADO), observaciones",
         "Eventos de mantenimiento por vehiculo."],
    ], widths_cm=[3.6, 7.2, 5.6])

    # 6.3 Modulo clientes/conductores
    add_heading(doc, "6.3. Modulo de clientes y conductores", level=2)
    add_table(doc, ["Tabla", "Campos clave", "Descripcion"], [
        ["clientes",
         "id_cliente (PK), codigo_cliente, tipo_identificacion, numero_identificacion, "
         "nombres, apellidos, fecha_nacimiento, telefono, correo, direccion_principal, estado",
         "Personas naturales/juridicas que rentan."],
        ["conductores",
         "id_conductor (PK), codigo, id_cliente (FK opcional), tipo_identificacion, "
         "numero_identificacion, nombres, apellidos, numero_licencia, fecha_vencimiento, edad, "
         "telefono, correo, es_conductor_joven (calc), estado",
         "Personas autorizadas a conducir el vehiculo."],
    ], widths_cm=[3.6, 7.2, 5.6])

    # 6.4 Modulo transaccional
    add_heading(doc, "6.4. Modulo transaccional (reservas, contratos, pagos, facturas)", level=2)
    add_table(doc, ["Tabla", "Campos clave", "Descripcion"], [
        ["reservas",
         "id_reserva (PK), reserva_guid, codigo_reserva, id_cliente (FK), id_vehiculo (FK), "
         "id_loc_recogida (FK), id_loc_devolucion (FK), canal, fecha_recogida, fecha_devolucion, "
         "subtotal, impuestos, extras, deposito, cargo_one_way, total, codigo_confirmacion, "
         "estado (PENDIENTE/CONFIRMADA/CANCELADA), requiere_hold, motivo_cancelacion, origen",
         "Reservas creadas por canal interno o externo."],
        ["reserva_extras",
         "id_reserva_extra (PK), id_reserva (FK), id_extra (FK), cantidad, valor_unitario, subtotal, estado",
         "Detalle de extras solicitados por reserva."],
        ["reserva_conductores",
         "id_reserva_conductor (PK), id_reserva (FK), id_conductor (FK), tipo, es_principal, cargo_conductor_joven, estado",
         "Conductores asignados a una reserva."],
        ["contratos",
         "id_contrato (PK), numero_contrato, id_reserva (FK), id_cliente (FK), id_vehiculo (FK), "
         "fecha_salida, fecha_prevista_devolucion, kilometraje_salida, nivel_combustible_salida, "
         "estado (ABIERTO/CERRADO), pdf_url, observaciones",
         "Apertura formal de la renta."],
        ["check_in_out",
         "id_check (PK), id_contrato (FK), tipo (IN/OUT), fecha_hora, kilometraje, "
         "nivel_combustible, limpio, observaciones, cargo_combustible, cargo_limpieza, cargo_km_extra",
         "Eventos de entrega y devolucion del vehiculo."],
        ["pagos",
         "id_pago (PK), codigo_pago, id_reserva (FK opc), id_contrato (FK opc), id_cliente (FK), "
         "tipo, metodo, estado, referencia_externa, monto, moneda, fecha_pago",
         "Transacciones monetarias."],
        ["facturas",
         "id_factura (PK), numero_factura, id_cliente (FK), id_reserva (FK opc), id_contrato (FK opc), "
         "fecha_emision, subtotal, iva, total, observaciones, estado, servicio_origen, motivo_inhab",
         "Comprobantes fiscales."],
    ], widths_cm=[3.6, 7.6, 5.2])

    # 6.5 Seguridad
    add_heading(doc, "6.5. Modulo de seguridad", level=2)
    add_table(doc, ["Tabla", "Campos clave", "Descripcion"], [
        ["security.usuarios_app",
         "id_usuario (PK), usuario_guid, username, correo, password_hash, password_salt, "
         "requiere_cambio_password, estado, activo, intentos_fallidos, bloqueado_hasta, "
         "ultimo_login, id_cliente (FK opc)",
         "Usuarios de la aplicacion."],
        ["security.roles",
         "id_rol (PK), rol_guid, nombre_rol, descripcion, es_sistema, estado",
         "Roles del sistema (ADMIN, AGENTE_POS, CLIENTE_WEB)."],
        ["security.usuarios_roles",
         "id_usuario_rol (PK), id_usuario (FK), id_rol (FK), estado, activo",
         "Asignacion N a M entre usuarios y roles."],
        ["security.permisos",
         "id_permiso (PK), codigo_permiso, modulo, accion, descripcion, estado",
         "Permisos granulares por modulo/accion."],
        ["security.roles_permisos",
         "id_rol_permiso (PK), id_rol (FK), id_permiso (FK), estado",
         "Permisos asignados a cada rol."],
        ["security.sesiones",
         "id_sesion (PK), sesion_guid, id_usuario (FK), token_id, refresh_token_hash, ip_origen, "
         "user_agent, fecha_inicio, fecha_expiracion, fecha_cierre, estado",
         "Sesiones JWT activas/cerradas."],
    ], widths_cm=[4.4, 7.2, 4.8])

    # 6.6 Auditoria
    add_heading(doc, "6.6. Modulo de auditoria", level=2)
    add_table(doc, ["Tabla", "Campos clave", "Descripcion"], [
        ["audit.aud_eventos",
         "id_aud_evento (PK), esquema_afectado, tabla_afectada, operacion (INSERT/UPDATE/DELETE), "
         "id_registro_afectado, datos_anteriores (JSON), datos_nuevos (JSON), usuario_app, "
         "login_bd, ip_origen, origen_evento, fecha_evento_utc",
         "Caja negra de cambios sobre tablas criticas."],
        ["audit.aud_intentos_login",
         "id_aud_login (PK), username_intentado, correo_intentado, resultado, motivo, ip_origen, "
         "user_agent, fecha_evento_utc",
         "Bitacora de intentos de inicio de sesion."],
    ], widths_cm=[4.4, 7.2, 4.8])

    add_heading(doc, "6.7. Relaciones clave", level=2)
    rel = [
        ["paises 1—N ciudades", "Cada ciudad pertenece a un pais."],
        ["ciudades 1—N localizaciones", "Las sucursales se asocian a una ciudad."],
        ["localizaciones 1—N vehiculos", "Cada vehiculo tiene una localizacion actual."],
        ["marcas/categorias 1—N vehiculos", "Catalogo aplicado a cada vehiculo."],
        ["clientes 1—N reservas", "Un cliente puede tener varias reservas."],
        ["reservas 1—1 contratos (opcional)", "Una reserva confirmada genera contrato."],
        ["contratos 1—N check_in_out", "Multiples eventos por contrato."],
        ["reservas 1—N pagos", "Una reserva puede tener varios pagos parciales."],
        ["clientes 1—N facturas", "Comprobantes asociados al cliente."],
        ["usuarios_app 1—N usuarios_roles N—1 roles", "RBAC."],
    ]
    add_table(doc, ["Relacion", "Descripcion"], rel, widths_cm=[6.0, 10.5])

    doc.add_page_break()

    # ============ 7. Backend y endpoints ============
    add_heading(doc, "7. Backend y endpoints de la API", level=1)
    add_paragraph(doc,
        "La API expone dos familias de endpoints. La primera, bajo /api/v1/admin y "
        "controladores RESTful (Clientes, Pagos, Contratos, Catalogos, Mantenimientos, etc.), "
        "esta orientada al back-office y al portal del cliente y requiere autenticacion JWT y "
        "autorizacion por rol. La segunda, bajo rutas publicas de /api/v1/ (vehiculos, "
        "localizaciones, reservas, etc.), implementa el contrato de integracion con OTA tipo Booking. "
        "Adicionalmente se exponen health checks y Swagger.")

    add_heading(doc, "7.1. Convenciones generales", level=2)
    add_bullet(doc, "Versionado: api/v{version}/... (actualmente 1.0).")
    add_bullet(doc, "Formato de respuesta: ApiResponse<T> { success, statusCode, message, data, traceId, errors }.")
    add_bullet(doc, "Autenticacion: JWT Bearer (header Authorization).")
    add_bullet(doc, "Autorizacion: por rol (ADMIN, AGENTE_POS, CLIENTE_WEB). Algunos endpoints permiten anonimo explicito.")
    add_bullet(doc, "Errores: middleware global devuelve respuestas tipificadas y traceId para soporte.")
    add_bullet(doc, "CORS: politica FrontendPolicy aplicada a toda la API.")
    add_bullet(doc, "Health checks: /health/live (proceso) y /health/ready (BD y dependencias).")

    # 7.2 Auth
    add_heading(doc, "7.2. AuthController", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_AUTH, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.3. Booking publico — Vehiculos", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_BOOKING_VEH, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.4. Booking publico — Catalogos", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_BOOKING_CAT, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.5. Booking publico — Reservas", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_BOOKING_RES, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.6. Admin — Vehiculos", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_VEHICULOS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.7. Admin — Localizaciones", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_LOCALIZ, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.8. Admin — Reservas", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_RESERVAS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.9. Admin — Contratos", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_CONTRATOS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.10. Admin — Pagos", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_PAGOS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.11. Admin — Facturas (cliente web)", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_FACTURAS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.12. Catalogos (autenticados)", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_CATALOGOS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.13. Admin — Clientes", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_CLIENTES, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.14. Admin — Usuarios", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_USUARIOS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.15. Admin — Mantenimientos", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_ADMIN_MANTENIMIENTOS, widths_cm=[2.0, 5.5, 6.5, 2.5])

    add_heading(doc, "7.16. Salud y documentacion", level=2)
    add_table(doc, ["Metodo", "Ruta", "Descripcion", "Autorizacion"], ENDPOINTS_HEALTH, widths_cm=[2.0, 5.5, 6.5, 2.5])

    doc.add_page_break()

    # ============ 8. Frontend ============
    add_heading(doc, "8. Construccion del frontend", level=1)
    add_paragraph(doc,
        "El frontend es una SPA construida con React 19 + Vite 8. Esta organizada por dominios "
        "funcionales: paginas publicas, panel administrativo y portal del cliente. La aplicacion "
        "consume la API mediante un cliente Axios centralizado, gestiona el estado del servidor "
        "con React Query y el estado global de UI/sesion con Zustand.")

    add_heading(doc, "8.1. Estructura de carpetas (frontend/src)", level=2)
    fe_struct = [
        ["api/", "Servicios HTTP por dominio (auth, booking, vehiculos, reservas, clientes, etc.) y axiosClient."],
        ["pages/", "Vistas agrupadas por modulo: home, auth, catalogo, reservar, dashboard, vehiculos, reservas, contratos, pagos, mantenimientos, usuarios, localizaciones, cliente."],
        ["components/", "Componentes reutilizables (layout, ui, ErrorBoundary)."],
        ["routes/", "ProtectedRoute (guardia por autenticacion y tipo de usuario)."],
        ["store/", "Stores Zustand (useAuthStore, useAppStore)."],
        ["utils/", "Helpers (errorHandler, validacion, manejo de reservas, integracion Cloudinary)."],
        ["main.jsx / App.jsx", "Bootstrap, configuracion de QueryClient, ruteo principal y Toaster global."],
    ]
    add_table(doc, ["Carpeta", "Contenido"], fe_struct, widths_cm=[5.0, 11.5])

    add_heading(doc, "8.2. Layouts y rutas", level=2)
    rutas = [
        ["/", "HomePage", "Inicio publico con buscador (Navbar)."],
        ["/buscar", "BuscarPage", "Resultados de disponibilidad."],
        ["/catalogo", "CatalogoPage", "Catalogo navegable de vehiculos."],
        ["/reservar/:id", "ReservarPage", "Flujo de reserva (cliente o invitado)."],
        ["/login", "LoginPage", "Inicio de sesion."],
        ["/registro", "RegisterPage", "Registro de cliente nuevo."],
        ["/dashboard", "DashboardPage", "Panel administrativo (ADMIN/AGENTE_POS) con MainLayout."],
        ["/clientes", "ClientesPage", "Gestion de clientes."],
        ["/vehiculos", "VehiculosPage", "Gestion de flota."],
        ["/reservas", "ReservasPage", "Gestion de reservas."],
        ["/contratos", "ContratosPage", "Gestion de contratos."],
        ["/pagos", "PagosPage", "Listado y registro de pagos."],
        ["/mantenimientos", "MantenimientosPage", "Mantenimientos por vehiculo."],
        ["/localizaciones", "LocalizacionesPage", "Sucursales."],
        ["/usuarios", "UsuariosPage", "Usuarios y roles."],
        ["/mi-cuenta", "MiCuentaPage", "Perfil del cliente (ClienteLayout)."],
        ["/mis-reservas", "MisReservasPage", "Reservas del cliente."],
        ["/mis-contratos", "MisContratosPage", "Contratos del cliente."],
        ["/mis-facturas", "MisFacturasPage", "Facturas del cliente."],
        ["/historial", "HistorialPage", "Historial general del cliente."],
        ["*", "NotFoundPage", "Fallback 404."],
    ]
    add_table(doc, ["Ruta", "Componente", "Descripcion"], rutas, widths_cm=[3.6, 4.0, 9.0])

    add_heading(doc, "8.3. Cliente HTTP y manejo de errores", level=2)
    add_paragraph(doc,
        "axiosClient.js centraliza el consumo de la API. Inyecta el token JWT en cada request, "
        "reintenta peticiones idempotentes ante errores transitorios (5xx/red), maneja 401 con "
        "limpieza de sesion y redireccion controlada (sin afectar rutas publicas), y muestra "
        "toasts de error mediante Sonner. La utilidad parseApiError normaliza la respuesta "
        "de error de la API para presentarla al usuario con un traceId opcional para soporte.")

    add_heading(doc, "8.4. Estado global", level=2)
    add_table(doc, ["Store", "Responsabilidad"], [
        ["useAuthStore",
         "Sesion (token, user, isAuthenticated, userType). Login, logout, helpers hasRole/hasAnyRole y derivacion de tipo de usuario."],
        ["useAppStore",
         "Estado de UI: visibilidad del sidebar y toggles."],
    ], widths_cm=[5.0, 11.5])

    add_heading(doc, "8.5. Servicios de API por dominio", level=2)
    add_table(doc, ["Archivo", "Endpoints que consume"], [
        ["authApi.js", "Login, registro, validaciones de cedula."],
        ["bookingApi.js", "Endpoints publicos de Booking (vehiculos, localizaciones, categorias, extras, reservas)."],
        ["vehiculosApi.js", "Listado, busqueda, CRUD interno de vehiculos."],
        ["reservasApi.js", "Reservas internas (creacion, confirmacion, cancelacion, consulta)."],
        ["contratosApi.js", "Contratos, check-out y check-in."],
        ["pagosApi.js", "Listado y registro de pagos."],
        ["facturasApi.js", "Facturas del cliente autenticado."],
        ["clientesApi.js", "Gestion de clientes."],
        ["conductoresApi.js", "Conductores asociados al cliente/reserva."],
        ["mantenimientosApi.js", "Mantenimientos de la flota."],
        ["localizacionesApi.js", "CRUD de sucursales."],
        ["catalogosApi.js", "Catalogos comunes (categorias, marcas, extras)."],
        ["usuariosApi.js", "Usuarios, estados y roles."],
    ], widths_cm=[4.0, 12.5])

    add_heading(doc, "8.6. Patrones de UI", level=2)
    add_bullet(doc, "Layouts: MainLayout (admin) y ClienteLayout (portal cliente), ambos con Navbar.")
    add_bullet(doc, "Componentes UI: ImageUploader (con Cloudinary), DateTimePicker, ErrorBoundary global.")
    add_bullet(doc, "Validacion: Zod + react-hook-form en formularios criticos (login, registro, reserva).")
    add_bullet(doc, "Notificaciones: Sonner con politica de no-spam para 401/422 y supresion configurable por mutacion/query.")
    add_bullet(doc, "QueryClient: retry inteligente (4xx no se reintenta), backoff exponencial y staleTime de 30s.")

    doc.add_page_break()

    # ============ 9. Seguridad ============
    add_heading(doc, "9. Seguridad, roles y autorizacion", level=1)
    add_paragraph(doc,
        "El acceso al sistema esta protegido por JWT con expiracion configurable (JwtSettings.ExpirationMinutes). "
        "Los controles se aplican via [Authorize] y [Authorize(Roles=...)] en los controladores. La autorizacion sigue un modelo RBAC con tres roles principales:")
    add_table(doc, ["Rol", "Alcance"], [
        ["ADMIN", "Control total: gestion de usuarios, flota, sucursales, reservas, contratos, pagos y mantenimientos."],
        ["AGENTE_POS", "Operacion en sucursal: gestion de flota, clientes, reservas, contratos, pagos y mantenimientos."],
        ["CLIENTE_WEB", "Portal del cliente: consulta de sus propias reservas, contratos y facturas; cancelacion de sus reservas."],
    ], widths_cm=[3.5, 13.0])
    add_paragraph(doc, "Reglas de autorizacion adicionales:")
    add_bullet(doc, "Un cliente solo puede ver/cancelar sus propias reservas (validacion explicita en ReservasController).")
    add_bullet(doc, "El cliente solo puede cancelar reservas con fecha de retiro futura.")
    add_bullet(doc, "Acciones criticas (DELETE de vehiculos, localizaciones y usuarios) requieren ADMIN.")
    add_bullet(doc, "Los endpoints publicos para Booking funcionan sin autenticacion pero estan acotados por contrato y reglas de negocio (lead time, prioridad de canal interno).")
    add_paragraph(doc,
        "Los eventos de seguridad (logins) y los cambios sobre tablas criticas se registran en audit.* "
        "para trazabilidad. La autenticacion incluye control de intentos fallidos y bloqueo temporal mediante "
        "campos en security.usuarios_app.")

    # ============ 10. Configuracion y despliegue ============
    add_heading(doc, "10. Configuracion, ejecucion y despliegue", level=1)
    add_heading(doc, "10.1. Requisitos", level=2)
    add_bullet(doc, ".NET SDK 10")
    add_bullet(doc, "Node.js 20+ y npm")
    add_bullet(doc, "PostgreSQL accesible (configurado para Supabase en appsettings.json)")

    add_heading(doc, "10.2. Configuracion del backend", level=2)
    add_paragraph(doc, "Archivo principal: src/Europcar.Rental.Api/appsettings.json. Variables clave:")
    add_bullet(doc, "ConnectionStrings:RentalDb")
    add_bullet(doc, "JwtSettings:SecretKey, Issuer, Audience, ExpirationMinutes")
    add_bullet(doc, "Logging: niveles globales y por categoria")
    add_paragraph(doc, "Recomendacion: mover secretos a variables de entorno o User Secrets (UserSecretsId 7f74ad0e-e0e5-4112-966e-57be085c77e1).")

    add_heading(doc, "10.3. Ejecucion local", level=2)
    add_paragraph(doc, "Backend (desde la raiz del repositorio):", bold=True)
    add_bullet(doc, "dotnet restore")
    add_bullet(doc, "dotnet run --project src/Europcar.Rental.Api")
    add_paragraph(doc, "URLs por defecto: http://localhost:5207 ; Swagger en /swagger ; Health: /health/live, /health/ready.")
    add_paragraph(doc, "Frontend (desde frontend/):", bold=True)
    add_bullet(doc, "npm ci")
    add_bullet(doc, "npm run dev")
    add_paragraph(doc, "VITE_API_URL debe apuntar al backend (ej. http://localhost:5207). El cliente Axios usa esta variable como baseURL.")

    add_heading(doc, "10.4. Despliegue del frontend (Render)", level=2)
    add_paragraph(doc, "El archivo render.yaml define el blueprint para publicar el frontend como sitio estatico:")
    add_bullet(doc, "Build: npm ci && npm run build")
    add_bullet(doc, "Publicacion: frontend/dist")
    add_bullet(doc, "Rewrite SPA: /* -> /index.html para soportar rutas del SPA")

    add_heading(doc, "10.5. Datos iniciales (seed)", level=2)
    add_paragraph(doc,
        "Al arrancar, la API ejecuta un seed resiliente que crea, si no existen, los usuarios "
        "admin.dev (ADMIN), agente.pos (AGENTE_POS) y cliente.web (CLIENTE_WEB), todos con contrasena por defecto 12345 (uso academico). "
        "Si falla, el log se reporta como warning y la API continua iniciando.")

    # ============ 11. RNF ============
    add_heading(doc, "11. Requerimientos no funcionales", level=1)
    add_table(doc, ["ID", "Requisito", "Medida concreta"], [
        ["RNF-01", "Disponibilidad", "99.9% mensual (max ~45 min de caida)."],
        ["RNF-02", "Recuperacion ante fallos", "RTO menor a 2 horas tras fallo de servidor de BD."],
        ["RNF-03", "Sincronizacion de reservas", "0% de perdida en transacciones confirmadas por canales externos."],
        ["RNF-04", "Latencia API", "Menos de 500 ms por peticion de disponibilidad bajo carga normal."],
        ["RNF-05", "Carga del frontend", "LCP menor a 2 segundos."],
        ["RNF-06", "Concurrencia BD", "Soporte minimo de 200 TPS sin deadlocks."],
        ["RNF-07", "Cifrado", "TLS 1.3 en transporte; AES-256 para datos sensibles en reposo."],
        ["RNF-08", "Seguridad de API", "JWT con expiracion y rotacion de claves; opcional OAuth2 para socios."],
        ["RNF-09", "Sesiones", "Cierre automatico tras 15 minutos de inactividad en el panel."],
        ["RNF-10", "Modularidad", "Logica en services/actions, no en controllers."],
        ["RNF-11", "Crecimiento de flota", "Soporte de hasta 50.000 vehiculos sin degradacion de busqueda."],
        ["RNF-12", "Compatibilidad de API", "Soporte simultaneo de versiones (v1, v2) para no romper integraciones."],
    ], widths_cm=[1.6, 4.4, 10.5])

    # ============ 12. Roadmap ============
    add_heading(doc, "12. Roadmap y mejoras propuestas", level=1)
    add_bullet(doc, "Agregar diagramas (arquitectura, ER, secuencia) embebidos como imagenes en docs/.")
    add_bullet(doc, "Exportar OpenAPI (swagger.json) por entorno y publicar coleccion Postman/Bruno.")
    add_bullet(doc, "Pipeline CI con pruebas unitarias e integracion (xUnit + Vitest/Playwright).")
    add_bullet(doc, "Migrar secretos a variables de entorno y User Secrets en local.")
    add_bullet(doc, "Refresh tokens con almacenamiento seguro y rotacion.")
    add_bullet(doc, "Cache de catalogos (categorias/extras/localizaciones) para mejorar latencia.")
    add_bullet(doc, "Trazabilidad con OpenTelemetry y correlationId end-to-end.")
    add_bullet(doc, "Hardening de auditoria: triggers a nivel de BD, no solo a nivel de aplicacion.")

    # ============ Anexo A. Diagramas funcionales ============
    doc.add_page_break()
    add_heading(doc, "Anexo A. Diagramas funcionales detallados", level=1)
    add_paragraph(doc,
        "Diagramas de siguiente nivel para cada modulo funcional, extraidos del documento de "
        "ingenieria del proyecto. Estos diagramas describen las acciones, entradas y salidas "
        "esperadas en cada modulo del sistema.")

    diagramas_funcionales = [
        ("figura_04.png", "Figura A1. Gestion de Vehiculos."),
        ("figura_05.png", "Figura A2. Gestion de Precios y Categorias."),
        ("figura_06.png", "Figura A3. Gestion de Locaciones."),
        ("figura_07.png", "Figura A4. Gestion de Contratos (Check-in / Check-out)."),
        ("figura_08.png", "Figura A5. Gestion de Pagos Internos."),
        ("figura_09.png", "Figura A6. Gestion de Clientes."),
        ("figura_10.png", "Figura A7. Reportes de Ocupacion."),
        ("figura_11.png", "Figura A8. Busqueda y Filtrado de Vehiculos."),
        ("figura_12.png", "Figura A9. Seleccion y Reserva de Vehiculos."),
        ("figura_13.png", "Figura A10. Pago en Linea."),
        ("figura_14.png", "Figura A11. Perfil y Historial de Usuario."),
        ("figura_15.png", "Figura A12. Consulta de Disponibilidad (API Booking)."),
        ("figura_16.png", "Figura A13. Sincronizacion de Reservas (API Booking)."),
    ]
    for filename, caption in diagramas_funcionales:
        add_image(doc, img(filename), caption=caption, width_cm=15.5)

    add_divider(doc)
    add_paragraph(doc, "Documento generado automaticamente a partir del codigo fuente y el contexto funcional del proyecto.",
                  italic=True, color=SOFT_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    try:
        doc.save(OUTPUT)
        print(f"Documento generado: {OUTPUT}")
    except PermissionError:
        fallback = OUTPUT.with_name(OUTPUT.stem + ".new" + OUTPUT.suffix)
        doc.save(fallback)
        print(
            "AVISO: el archivo original esta abierto en Word.\n"
            f"Se guardo una copia nueva en: {fallback}\n"
            "Cierra Word y vuelve a ejecutar para sobreescribir el original."
        )


if __name__ == "__main__":
    main()
