from docx import Document

src_path = r"Contexto/EUROPCAR.docx"
tgt_path = r"Contexto/Documentacion_Servicio_Europcar_V1.docx"

src = Document(src_path)
doc = Document(tgt_path)

rf_rows = []
for table in src.tables:
    if not table.rows:
        continue
    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    joined = " | ".join(headers)
    if "id" in joined and ("requerimiento" in joined or "requisito" in joined) and ("prioridad" in joined or "descripci" in joined):
        for r in table.rows[1:]:
            vals = [c.text.strip() for c in r.cells]
            if len(vals) >= 4 and vals[0].upper().startswith("F"):
                rf_rows.append(vals[:4])

seen = set()
rf_unique = []
for row in rf_rows:
    rid = row[0].strip().upper()
    if rid and rid not in seen:
        seen.add(rid)
        rf_unique.append(row)

if not rf_unique:
    raise RuntimeError("No se encontraron requerimientos funcionales (F*) en EUROPCAR.docx")

anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith("5.1."):
        anchor = p
        break

if anchor is None:
    raise RuntimeError("No se encontró el punto 5.1 en el documento objetivo")

# Insert 5.2 heading/description before old 5.1
p_52 = anchor.insert_paragraph_before("5.2. Requerimientos funcionales")
p_52.style = anchor.style
p_52_desc = anchor.insert_paragraph_before(
    "Tabla de requerimientos funcionales consolidada a partir del documento base EUROPCAR.docx."
)

# Add and reposition table after 5.2 description
tbl = doc.add_table(rows=len(rf_unique) + 1, cols=4)
tbl.style = "Table Grid"
headers = ["ID", "Requerimiento", "Descripción", "Prioridad"]
for i, h in enumerate(headers):
    tbl.cell(0, i).text = h

for r_idx, row in enumerate(rf_unique, start=1):
    for c_idx in range(4):
        tbl.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""

p_52_desc._p.addnext(tbl._tbl)
anchor.insert_paragraph_before("")

# Insert 5.1 heading and business rules
p_51 = p_52.insert_paragraph_before("5.1. Reglas del negocio")
p_51.style = anchor.style

rules = [
    "RN-01: No se puede confirmar una reserva con fecha de devolución menor o igual a la fecha de recogida.",
    "RN-02: No se permite reservar ni confirmar vehículos en estado operativo no disponible (mantenimiento, taller, alquilado o fuera de servicio).",
    "RN-03: No se permiten solapamientos de reservas activas para el mismo vehículo en el mismo rango de fechas.",
    "RN-04: La cancelación de reservas finalizadas o ya canceladas está bloqueada por reglas de negocio.",
    "RN-05: La cancelación de reservas por cliente solo se permite antes de la fecha de recogida.",
    "RN-06: El total de la reserva debe corresponder a subtotal + impuestos + extras + cargo one-way.",
    "RN-07: La confirmación de reserva debe generar pago y factura asociados de forma transaccional.",
    "RN-08: El extra de conductor adicional solo aplica cuando existe más de un conductor activo en la reserva.",
    "RN-09: El estado operativo ALQUILADO del vehículo se controla por flujos operativos y no por edición manual directa.",
    "RN-10: Toda operación sensible debe registrar trazabilidad de usuario y fecha de modificación para auditoría.",
]

for rule in reversed(rules):
    pr = p_52.insert_paragraph_before(rule)
    pr.style = p_52_desc.style

# Renumber existing 5.x points
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("5.1. Diagrama conceptual"):
        p.text = t.replace("5.1.", "5.3.", 1)
    elif t.startswith("5.2. Diagrama general de funciones"):
        p.text = t.replace("5.2.", "5.4.", 1)
    elif t.startswith("5.3. Diagramas de secuencia de las actividades principales"):
        p.text = t.replace("5.3.", "5.5.", 1)
    elif t.startswith("5.4. Detalle de casos de uso"):
        p.text = t.replace("5.4.", "5.6.", 1)

try:
    doc.save(tgt_path)
    print(f"Documento actualizado correctamente con {len(rf_unique)} requerimientos funcionales: {tgt_path}")
except PermissionError:
    alt_path = r"Contexto/Documentacion_Servicio_Europcar_V1.actualizado.docx"
    doc.save(alt_path)
    print(f"Archivo bloqueado; se guardó versión actualizada en: {alt_path}")
