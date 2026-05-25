# -*- coding: utf-8 -*-
"""Reconstruye Laboratorio_Gestion_Secretos(1).docx con numeración coherente y secciones anidadas."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(r"c:\Users\medin\source\repos\Proyecto Desarrollo")
DOC_PATH = ROOT / "Contexto" / "Laboratorio_Gestion_Secretos(1).docx"
BACKUP = DOC_PATH.with_suffix(".docx.bak")
IMG_DIR = ROOT / "Contexto" / "imagenes_laboratorio"

# Extraídas del docx (o colocar capturas con estos nombres)
IMG_CLIENTE_ENDPOINT = IMG_DIR / "image2.png"
IMG_CLIENTE_RESPUESTA = IMG_DIR / "image3.png"
IMG_CLIENTE_CAMPOS = IMG_DIR / "image4.png"


class Counters:
    def __init__(self) -> None:
        self.table = 0
        self.figure = 0

    def next_table(self, title: str) -> str:
        self.table += 1
        return f"Tabla {self.table}. {title}"

    def next_figure(self, title: str) -> str:
        self.figure += 1
        return f"Figura {self.figure}. {title}"


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def clear_body_after_cover(doc: Document, keep_paragraphs: int = 10) -> None:
    for p in list(doc.paragraphs[keep_paragraphs:]):
        delete_paragraph(p)
    body = doc.element.body
    for tbl in list(doc.tables):
        body.remove(tbl._tbl)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_normal(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Normal")


def add_figure_caption(doc: Document, counters: Counters, title: str) -> str:
    caption = counters.next_figure(title)
    add_normal(doc, caption)
    return caption


def add_code_figure(doc: Document, counters: Counters, title: str, code: str) -> None:
    add_figure_caption(doc, counters, title)
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


def add_image_figure(
    doc: Document,
    counters: Counters,
    title: str,
    image_path: Path,
    width_inches: float = 6.0,
) -> None:
    add_figure_caption(doc, counters, title)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width_inches))
        doc.add_paragraph()
    else:
        add_normal(doc, f"[Colocar captura: {image_path.name}]")


def build_content(doc: Document) -> None:
    c = Counters()

    add_heading(doc, "Introducción", 1)
    add_normal(
        doc,
        "La presente práctica aplica buenas prácticas de ciberseguridad en Europcar Rental "
        "(API .NET, frontend React/Vite y PostgreSQL). Se documentan hallazgos desde la perspectiva "
        "de un atacante externo, evidencia en código y en el navegador, las soluciones incorporadas "
        "en el repositorio y la validación de riesgos mitigados.",
    )

    add_heading(doc, "Objetivo del laboratorio", 1)
    add_normal(
        doc,
        "Aplicar manejo seguro de secretos y configuración por ambiente, identificar vulnerabilidades, "
        "corregirlas en el código y validar que la aplicación funcione sin exponer credenciales ni "
        "datos personales innecesarios en respuestas públicas.",
    )

    add_normal(doc, c.next_table("Herramientas requeridas y uso dentro del proyecto"))
    add_table(
        doc,
        ["Herramienta", "Uso en el proyecto"],
        [
            ["Visual Studio Code", "Edición del código y DevTools del navegador."],
            ["Git y GitHub", "Control de versiones; .gitignore para .env."],
            ["Node.js", "Frontend React/Vite en frontend/."],
            [".env / User Secrets", "Secretos fuera del repositorio."],
            ["PostgreSQL / Supabase", "Base de datos vía ConnectionStrings:RentalDb."],
            ["PowerShell / dotnet", "Variables de entorno y ejecución de la API."],
        ],
    )

    # --- FASE 1 ---
    add_normal(doc, "Fase 1 — Investigación inicial")
    add_normal(
        doc,
        "Se definió el alcance: monolito legado (_legacy/EuropcarRental) para administración y autenticación, "
        "middleware Middleware.RedCar para el contrato Booking V2, y cliente web React.",
    )
    add_normal(doc, c.next_table("Decisiones técnicas documentadas para el repositorio"))
    add_table(
        doc,
        ["Elemento", "Elección"],
        [
            ["API legada", "_legacy/EuropcarRental/src/Europcar.Rental.Api"],
            ["API orquestador", "Middleware.RedCar/src/Middleware.RedCar.Api"],
            ["Cliente web", "frontend/ (React + Vite)"],
            ["Autenticación", "JWT en cookie HttpOnly rc_auth (post-corrección)"],
            ["Secretos .NET", "ConnectionStrings__RentalDb, JwtSettings__SecretKey"],
            ["Secretos frontend", "VITE_* en .env local (no versionado)"],
        ],
    )

    # --- FASE 2 ---
    add_normal(doc, "Fase 2 — Hallazgos iniciales (antes del endurecimiento)")
    add_normal(
        doc,
        "Los riesgos se analizaron como los explotaría un atacante sin acceso al servidor: abuso de "
        "endpoints HTTP, inspección de respuestas en el navegador y filtración de artefactos (Git, backups).",
    )
    add_normal(doc, c.next_table("Resumen de riesgos identificados"))
    add_table(
        doc,
        ["Id", "Riesgo", "Severidad"],
        [
            ["A", "Catálogo accesible bajo ruta admin sin JWT", "Alta / media"],
            ["B", "guest-client devolvía PII al confirmar cédula", "Alta"],
            ["C", "JWT persistido en localStorage", "Alta (con XSS)"],
            ["D", "Swagger publicado sin restricción de entorno", "Media"],
            ["E", "CORS con métodos y cabeceras demasiado amplios", "Baja / media"],
            ["F", "SecretKey y cadena BD en appsettings.json", "Alta"],
            ["G", "Contraseña embebida en scripts/KillZombies.cs", "Alta"],
        ],
    )

    add_heading(doc, "2.1 Evidencia en código (estado inicial)", 2)
    add_code_figure(
        doc,
        c,
        "Controlador de vehículos — AllowAnonymous bajo ruta admin",
        '[HttpGet("disponibles")]\n[AllowAnonymous]  // corregido: [Authorize(Roles = "ADMIN,AGENTE_POS")]',
    )
    add_code_figure(
        doc,
        c,
        "Endpoint guest-client — respuesta con datos personales",
        "// Antes: IdCliente, Nombre1, Apellido1, NumeroIdentificacion, Correo\n"
        "// Ahora: solo { idCliente, esNuevo }",
    )
    add_code_figure(
        doc,
        c,
        "Frontend — token en localStorage",
        "// Antes: localStorage.setItem('token', loginResponse.token)\n"
        "// Ahora: cookie HttpOnly rc_auth + sessionStorage solo para perfil",
    )

    add_heading(doc, "2.2 Evidencia desde el lado del cliente (estado inicial)", 2)
    add_normal(
        doc,
        "Las capturas siguientes se obtuvieron con las herramientas de desarrollador del navegador "
        "(pestaña Red) durante el flujo de catálogo, antes de aplicar las correcciones en backend y frontend. "
        "Demuestran que un usuario sin rol administrativo podía observar rutas con prefijo admin y "
        "respuestas con campos operativos de más.",
    )
    add_image_figure(
        doc,
        c,
        "Solicitud HTTP visible en el navegador (ruta con prefijo admin)",
        IMG_CLIENTE_ENDPOINT,
        6.2,
    )
    add_normal(
        doc,
        "Se observa una petición hacia la API desplegada (Render) consumiendo "
        "/api/v1/admin/Vehiculos/disponibles desde la vista de catálogo. Las rutas administrativas "
        "no deben usarse como API pública del cliente.",
    )
    add_image_figure(
        doc,
        c,
        "Respuesta HTTP 200 con arreglo data al navegador",
        IMG_CLIENTE_RESPUESTA,
        6.2,
    )
    add_normal(
        doc,
        "La respuesta confirma que el listado completo de vehículos llega al navegador. "
        "Tras la corrección, el catálogo público usa GET /api/v1/vehiculos (Booking) y el listado "
        "admin exige JWT con rol ADMIN o AGENTE_POS.",
    )
    add_image_figure(
        doc,
        c,
        "Campos operativos visibles en el JSON (data)",
        IMG_CLIENTE_CAMPOS,
        6.2,
    )
    add_normal(
        doc,
        "Campos como idVehiculo, vehiculoGuid, codigoInterno o placa son de gestión interna. "
        "La mitigación incluye DTOs reducidos en la API pública y separación estricta de rutas "
        "admin frente a rutas de Booking.",
    )

    # --- FASE 3 — SOLUCIONES DETALLADAS ---
    add_normal(doc, "Fase 3 — Corrección segura y endurecimiento (soluciones implementadas)")
    add_normal(
        doc,
        "A continuación se detalla cada mitigación aplicada en el repositorio. La Tabla 4 resume el "
        "cumplimiento de requisitos del laboratorio.",
    )

    add_heading(doc, "3.1 Secretos fuera del código y del repositorio", 2)
    add_normal(
        doc,
        "Se vaciaron ConnectionStrings:RentalDb y JwtSettings:SecretKey en appsettings.json del monolito. "
        "La configuración real se carga por variables de entorno (ConnectionStrings__RentalDb, "
        "JwtSettings__SecretKey), User Secrets en desarrollo o variables del host en producción. "
        "Plantilla: _legacy/EuropcarRental/.env.example.",
    )
    add_code_figure(
        doc,
        c,
        "appsettings.json sin valores secretos (repositorio actual)",
        '{\n  "ConnectionStrings": { "RentalDb": "" },\n'
        '  "JwtSettings": { "SecretKey": "", "Issuer": "Europcar.Rental.Api", ... }\n}',
    )
    add_normal(
        doc,
        "El script scripts/KillZombies.cs dejó de incluir contraseñas en texto plano y lee la misma "
        "cadena que la API desde el entorno.",
    )
    add_code_figure(
        doc,
        c,
        "KillZombies.cs — lectura segura de la cadena",
        "var cs = Environment.GetEnvironmentVariable(\"ConnectionStrings__RentalDb\")\n"
        "    ?? Environment.GetEnvironmentVariable(\"RENTAL_DB_CONNECTION\")\n"
        "    ?? throw new InvalidOperationException(...);",
    )
    add_code_figure(
        doc,
        c,
        "Ejecución del script con variable de entorno (PowerShell)",
        '$env:ConnectionStrings__RentalDb = "Host=...;Password=...;..."\n'
        'dotnet script "scripts\\KillZombies.cs"',
    )

    add_heading(doc, "3.2 Mitigación A — Rutas admin protegidas", 2)
    add_normal(
        doc,
        "Se eliminó [AllowAnonymous] en GET admin/vehiculos/disponibles y GET admin/vehiculos/{id}. "
        "Ambos endpoints exigen rol ADMIN o AGENTE_POS. El catálogo y la reserva del canal público "
        "utilizan BookingVehiculosController bajo /api/v1/vehiculos, alineado al contrato RedCar.",
    )

    add_heading(doc, "3.3 Mitigación B — guest-client sin filtración de PII", 2)
    add_normal(
        doc,
        "El flujo de cliente invitado se movió a POST /api/v1/reservas/guest-client (ruta pública de reservas). "
        "La respuesta solo incluye idCliente y esNuevo, sin nombre, correo ni documento cuando el cliente "
        "ya existía. Se añadió rate limiting (20 solicitudes por minuto por IP) en guest-client y "
        "cedula-exists. El middleware expone el mismo contrato en LegacyPublicReservasController.",
    )

    add_heading(doc, "3.4 Mitigación C — Sesión en cookie HttpOnly", 2)
    add_normal(
        doc,
        "El login y el registro establecen la cookie HttpOnly rc_auth con el JWT. El cuerpo JSON ya no "
        "incluye el campo token. El frontend usa axios con withCredentials: true; el perfil de usuario "
        "se guarda en sessionStorage (no el token). GET /Auth/me restaura la sesión al recargar la página. "
        "POST /Auth/logout elimina la cookie.",
    )

    add_heading(doc, "3.5 Mitigación D — Swagger condicionado", 2)
    add_normal(
        doc,
        "En el monolito, UseSwagger solo se ejecuta si el entorno es Development o si "
        "Swagger:Enabled es true en configuración. En producción por defecto no se publica la "
        "documentación interactiva, reduciendo el reconocimiento automático de la superficie.",
    )

    add_heading(doc, "3.6 Mitigación E — CORS acotado", 2)
    add_normal(
        doc,
        "La política FrontendPolicy mantiene orígenes explícitos (localhost y despliegues del frontend), "
        "habilita AllowCredentials para la cookie de sesión, y limita métodos HTTP y cabeceras "
        "aceptadas en lugar de AllowAnyMethod/AllowAnyHeader.",
    )

    add_heading(doc, "3.7 Seed de usuarios solo en desarrollo", 2)
    add_normal(
        doc,
        "El seed de usuarios demo (admin.dev, agente.pos, cliente.web) solo corre cuando "
        "ASPNETCORE_ENVIRONMENT es Development. La contraseña de seed se lee de SEED_DEV_PASSWORD "
        "con valor por defecto documentado solo para laboratorio local.",
    )

    add_normal(doc, c.next_table("Estado de endurecimiento del repositorio"))
    add_table(
        doc,
        ["Requisito", "Estado", "Evidencia"],
        [
            [".env frontend", "Hecho", "frontend/.env.example; .env en .gitignore"],
            [".gitignore", "Hecho", "Entrada .env en raíz"],
            ["Script KillZombies", "Hecho", "Variables de entorno"],
            ["appsettings sin secretos", "Hecho", "Valores vacíos + .env.example"],
            ["Superficie HTTP A–E", "Hecho", "Secciones 3.2–3.6"],
            ["Separación dev/prod", "Hecho", "User Secrets; seed solo Development"],
            ["Mínimo privilegio BD (V2)", "Parcial", "db/microservices/99_supabase_grants.sql (roles ms_*)"],
        ],
    )

    # --- FASE 4 ---
    add_normal(doc, "Fase 4 — Validación, análisis y pruebas")
    add_normal(doc, c.next_table("Lista de comprobación para la validación"))
    add_table(
        doc,
        ["Prueba", "Resultado esperado"],
        [
            ["API con secretos por entorno", "Arranque sin contraseñas en archivos trackeados"],
            ["Login", "Cookie rc_auth; JSON sin token"],
            ["guest-client", "Solo idCliente y esNuevo"],
            ["Admin sin JWT", "GET admin/vehiculos/disponibles → 401"],
            ["Catálogo cliente", "Consumo de /api/v1/vehiculos, no /admin/..."],
            ["Swagger producción", "Deshabilitado salvo Swagger__Enabled=true"],
            ["Capturas", "Sin contraseñas ni JWT completos visibles"],
        ],
    )

    # --- ENTREGABLES ---
    add_heading(doc, "Resultados esperados del laboratorio", 1)
    add_normal(doc, c.next_table("Entregables y forma de demostración"))
    add_table(
        doc,
        ["Entregable", "Demostración"],
        [
            ["Aplicación funcional y segura", "Flujos login, catálogo y reserva invitado tras correcciones"],
            ["Variables protegidas", ".env ignorado; secretos por entorno"],
            ["Código corregido", "Ramas con mitigaciones A–G y frontend alineado"],
            ["Evidencia cliente", "Figuras 4–6 (antes) y validación post-corrección en Fase 4"],
            ["Informe técnico", "Este documento"],
        ],
    )

    add_heading(doc, "Conclusiones", 1)
    for text in [
        "Los secretos no deben versionarse: appsettings.json y scripts quedaron sin credenciales reales.",
        "La seguridad HTTP requiere rutas públicas dedicadas, respuestas mínimas y rate limiting en endpoints abusables.",
        "El navegador es parte del perímetro: las Figuras 4–6 muestran por qué se separaron rutas admin y se adoptó cookie HttpOnly.",
        "Swagger y CORS deben configurarse de forma restrictiva en producción.",
        "User Secrets y variables de entorno permiten separar ambientes; rotar credenciales si alguna vez se expusieron en Git.",
    ]:
        add_normal(doc, text)


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"No existe: {DOC_PATH}")

    if not BACKUP.exists():
        shutil.copy2(DOC_PATH, BACKUP)

    doc = Document(str(DOC_PATH))
    clear_body_after_cover(doc, keep_paragraphs=10)
    build_content(doc)
    doc.save(str(DOC_PATH))
    print(f"Documento actualizado: {DOC_PATH}")
    print(f"Figuras/tablas renumeradas; imágenes desde: {IMG_DIR}")


if __name__ == "__main__":
    main()
