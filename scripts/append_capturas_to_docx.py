# -*- coding: utf-8 -*-
"""Añade capturas de validación (Fase 4) al docx del laboratorio."""
from pathlib import Path

from docx import Document
from docx.shared import Inches

ROOT = Path(r"c:\Users\medin\source\repos\Proyecto Desarrollo")
DOC = ROOT / "Contexto" / "Laboratorio_Gestion_Secretos(1).docx"
DOC_OUT = ROOT / "Contexto" / "Laboratorio_Gestion_Secretos_CON_CAPTURAS.docx"
IMG = ROOT / "Contexto" / "imagenes_laboratorio" / "informe"

# Continúa numeración tras Figura 9 del documento actual
FIG_START = 10

CAPTIONS = [
    ("10_swagger_monolito_inicio.png", "Swagger UI del monolito (entorno Development)"),
    ("11_swagger_monolito_auth.png", "Swagger — sección Auth"),
    ("12_swagger_monolito_booking.png", "Swagger — endpoints Booking / públicos"),
    ("13_login_sin_token_en_json.png", "Login admin.dev: JSON sin campo token (cookie HttpOnly)"),
    ("14_guest_client_respuesta_minima.png", "POST guest-client: respuesta solo idCliente y esNuevo"),
    ("16_validacion_cookie_y_401_admin.png", "Validación: cookie rc_auth y HTTP 401 en admin sin JWT"),
]


def insert_before(paragraph, text: str, style: str = "Normal"):
    new_p = paragraph.insert_paragraph_before(text, style)
    return new_p


def main() -> None:
    doc = Document(str(DOC))
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("Resultados esperados del laboratorio"):
            anchor = p
            break
    if anchor is None:
        raise SystemExit("No se encontró ancla 'Resultados esperados'")

    insert_before(anchor, "", "Normal")
    insert_before(anchor, "4.1 Evidencia de validación en ejecución local", "Heading 2")
    insert_before(
        anchor,
        "API en http://localhost:5207 (Development). Health /health/ready: Healthy. "
        "Swagger disponible en /swagger. Las capturas siguientes documentan el comportamiento tras las correcciones.",
        "Normal",
    )

    blocks = [
        (FIG_START + i, filename, title)
        for i, (filename, title) in enumerate(CAPTIONS)
    ]

    for num, filename, title in reversed(blocks):
        insert_before(anchor, f"Figura {num}. {title}", "Normal")
        path = IMG / filename
        if path.exists():
            p = anchor.insert_paragraph_before("")
            run = p.add_run()
            run.add_picture(str(path), width=Inches(6.2))
            anchor.insert_paragraph_before("", "Normal")
        else:
            insert_before(anchor, f"[Falta archivo: {filename}]", "Normal")

    out = DOC_OUT
    try:
        doc.save(str(DOC))
        out = DOC
    except PermissionError:
        doc.save(str(DOC_OUT))
        out = DOC_OUT
    print(f"Insertadas figuras {FIG_START}–{FIG_START + len(CAPTIONS) - 1} en {out}")


if __name__ == "__main__":
    main()
