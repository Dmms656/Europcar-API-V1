# -*- coding: utf-8 -*-
"""Convierte docs/Laboratorio-Gestion-Secretos-y-Configuracion-Segura.md a .docx."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def add_body_paragraph(doc: Document, text: str) -> None:
    if not text.strip():
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    add_paragraph_runs_from_markdown(p, text)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def add_code_block(doc: Document, lines: list[str]) -> None:
    text = "\n".join(lines)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def main() -> int:
    root = Path(__file__).resolve().parent.parent
    md_path = root / "docs" / "Laboratorio-Gestion-Secretos-y-Configuracion-Segura.md"
    out_path = root / "docs" / "Laboratorio-Gestion-Secretos-y-Configuracion-Segura.docx"

    if not md_path.is_file():
        print(f"No se encuentra: {md_path}", file=sys.stderr)
        return 1

    raw_lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(raw_lines):
        line = raw_lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if is_table_row(line):
            rows: list[list[str]] = []
            while i < len(raw_lines) and is_table_row(raw_lines[i]):
                rows.append(parse_table_row(raw_lines[i]))
                i += 1
            if len(rows) >= 2 and is_separator_row(rows[1]):
                rows.pop(1)
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row_cells in enumerate(rows):
                    for ci in range(ncols):
                        val = row_cells[ci] if ci < len(row_cells) else ""
                        table.rows[ri].cells[ci].text = val
                doc.add_paragraph()
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("# "):
            t = doc.add_heading(line[2:].strip(), level=0)
            t.runs[0].font.size = Pt(16)
            i += 1
            continue

        if line.strip().startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            add_paragraph_runs_from_markdown(p, line.strip()[6:].strip())
            i += 1
            continue
        if line.strip().startswith("- "):
            body = line.strip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_paragraph_runs_from_markdown(p, body)
            i += 1
            continue

        if line.strip().startswith("> "):
            add_body_paragraph(doc, line.strip()[2:])
            i += 1
            continue

        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        add_body_paragraph(doc, line)
        i += 1

    try:
        doc.save(out_path)
        print(f"Generado: {out_path}")
    except PermissionError:
        alt = out_path.with_name(out_path.stem + "-export.docx")
        doc.save(alt)
        print(f"No se pudo sobrescribir (archivo abierto). Generado: {alt}")
    return 0


def add_paragraph_runs_from_markdown(p, text: str) -> None:
    p.clear()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            p.add_run(part)


if __name__ == "__main__":
    raise SystemExit(main())
